# -*- coding: utf-8 -*-
"""Small regression tests for the clean-baseline integrity audit."""

from __future__ import annotations

import copy
import sys
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

from audit_review_integrity import audit

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def W(name: str) -> str:
    return f"{{{W_NS}}}{name}"


def revise_paragraph(root, index: int, replacement: str) -> None:
    paragraph = root.xpath("//w:body//w:p", namespaces=NS)[index]
    run = paragraph.find(W("r"))
    old_run = copy.deepcopy(run)
    old_text = old_run.find(W("t"))
    old_text.tag = W("delText")
    deleted = etree.Element(W("del"), {W("id"): "1", W("author"): "Test", W("date"): "2026-01-01T00:00:00Z"})
    deleted.append(old_run)
    inserted = etree.Element(W("ins"), {W("id"): "2", W("author"): "Test", W("date"): "2026-01-01T00:00:00Z"})
    new_run = copy.deepcopy(run)
    new_run.find(W("t")).text = replacement
    inserted.append(new_run)
    run.getparent().replace(run, deleted)
    deleted.addnext(inserted)


def make_revised(baseline: Path, output: Path, revised_indexes: dict[int, str]) -> None:
    with zipfile.ZipFile(baseline) as source:
        document = etree.fromstring(source.read("word/document.xml"))
        for index, text in revised_indexes.items():
            revise_paragraph(document, index, text)
        with zipfile.ZipFile(output, "w") as target:
            for item in source.infolist():
                data = etree.tostring(document, xml_declaration=True, encoding="UTF-8", standalone="yes") if item.filename == "word/document.xml" else source.read(item.filename)
                target.writestr(item, data)


def main() -> int:
    with tempfile.TemporaryDirectory(prefix="docx_audit_test_") as temp:
        root = Path(temp)
        baseline = root / "baseline.docx"
        revised = root / "revised.docx"
        bad = root / "bad.docx"
        doc = Document()
        doc.add_paragraph("Keep unchanged")
        doc.add_paragraph("Old word")
        doc.save(baseline)

        make_revised(baseline, revised, {1: "New word"})
        errors, _, _ = audit(baseline, revised, {1}, False, False)
        assert not errors, errors

        make_revised(baseline, bad, {0: "Unexpected change", 1: "New word"})
        errors, _, _ = audit(baseline, bad, {1}, False, False)
        assert any("non-target paragraph 0" in error for error in errors), errors
    print("AUDIT TESTS OK")
    return 0


if __name__ == "__main__":
    sys.exit(main())
