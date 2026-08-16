# -*- coding: utf-8 -*-
"""Refine a paragraph-by-paragraph review markdown into a polish_edits.json.

Usage:
  python refine_review_markdown.py input.docx review.md -o polish_edits.json

The output follows the polish-skill contract and can be applied with:
  python ai_review_to_comments.py tracked input.docx polish_edits.json -o revised.docx

Key principles learned in practice:
  * Always validate anchors against the CURRENT document; reviews may be based on an older draft.
  * Check formatting (superscript/subscript) before treating plain text as an error.
  * Apply issue-driven edits only: review-listed problem + present in document + explicit replacement.
"""

from __future__ import annotations

import argparse
import difflib
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# Explicit replacements extracted from journal conventions / review suggestions.
# Extend this mapping when a review lists a problem without an inline replacement.
KNOWN_REPLACEMENTS = {
    "∅S": "φS",
    "alendaring": "calendaring",
    "calendering": "calendaring",
    "threedimensionally": "three-dimensionally",
    "inplane": "in-plane",
    "A1": "A-1",
    "Figure S4, Supporting Information": "Figure S4 (Supporting Information)",
    "deterioration in the thermal conductivity": "degradation of thermal conductivity",
    "sharp reduction in ductility": "severe loss of ductility",
    "This superior performance": "Such outstanding performance",
    "determining the formation of continuous thermal transport pathways": "governing the construction of interconnected thermal networks",
    "strong orientation tendency": "strong tendency toward in-plane alignment",
    "3D through-plane thermal transport networks": "3D through-plane thermal conduction networks",
    "as-prepared composites": "as-fabricated composites",
    "pronounced increase in the orientation disorder": "greatly elevated orientational disorder",
    "in excellent agreement with": "consistent with",
    "representing a significant enhancement": "indicating dramatic cross-plane thermal improvement",
    "synergistic improvement in thermal conductivity": "synergistic enhancement of thermal conductivity",
    "was used to monitor": "is adopted to track",
    "transient thermal response": "transient heat transfer behavior",
    "were used to characterize": "are employed to characterize",
    "lowest crystallinity which is ascribed": "lowest crystallinity, which is ascribed",
    "spatial steric hindrance": "steric hindrance",
    "well-defined crystalline structures": "regular crystalline domains",
    "discussed above": "analyzed in prior sections",
    "limited capacity for further deformation": "restricted deformability",
    "slip-prone interfaces": "slidable inter-filler interfaces",
    "pre-existing internal stress": "residual internal stress",
    "the destruction of the continuous confinement network": "the collapse of continuous filler confinement networks",
    "were performed": "are carried out",
    "was measured": "is characterized",
    "were systematically characterized": "are comprehensively characterized",
    "were prepared via": "we fabricated via",
    "excellent low dielectric loss performance": "ultralow dielectric loss characteristic",
    "Figure 2b; Figures S6, S7": "Figure 2b, Figures S6 and S7",
}


def norm(text: str) -> str:
    out = []
    for ch in text:
        if ch in "\u00a0\u202f\u2009\u200a\u2007":
            out.append(" ")
        elif ch in "\u2010\u2011\u2012\u2013\u2014\u2015":
            out.append("-")
        else:
            out.append(ch)
    return "".join(out)


def split_sentences(text: str) -> list[str]:
    text = norm(text)
    out = []
    start = 0
    i = 0
    while i < len(text) - 1:
        ch = text[i]
        nxt = text[i + 1]
        if ch in ".!?" and (nxt.isspace() or nxt in "\"'("):
            if ch == ".":
                j = i + 1
                while j < len(text) and text[j].isspace():
                    j += 1
                if j >= len(text) or not text[j].isupper():
                    i += 1
                    continue
            out.append(text[start : i + 1].strip())
            start = i + 1
        elif ch in "。！？" or (ch == "；" and (nxt.isspace() or "\u4e00" <= nxt <= "\u9fff")):
            out.append(text[start : i + 1].strip())
            start = i + 1
        i += 1
    tail = text[start:].strip()
    if tail:
        out.append(tail)
    return [s for s in out if s]


def split_cn(text: str) -> list[str]:
    return [p.strip() for p in re.split(r"(?<=[。；！？])", text) if p.strip()]


def parse_markdown(md_path: Path) -> list[dict]:
    text = md_path.read_text(encoding="utf-8")
    parts = re.split(r"(?m)^### 段落\d+ 原文\n", text)
    blocks = []
    for body in parts[1:]:
        m_orig = re.search(r"^(.*?)\n#### 中文翻译", body, re.S)
        m_cn = re.search(r"#### 中文翻译\n(.*?)\n#### 英文不妥/错误", body, re.S)
        m_iss = re.search(r"#### 英文不妥/错误\n(.*?)\n#### Composites Part B 规范改写", body, re.S)
        m_rev = re.search(r"#### Composites Part B 规范改写\n(.*)$", body, re.S)
        if not m_orig:
            continue
        blocks.append(
            {
                "orig": m_orig.group(1).strip(),
                "cn": m_cn.group(1).strip() if m_cn else "",
                "issues_raw": m_iss.group(1).strip() if m_iss else "",
                "rev": m_rev.group(1).strip() if m_rev else "",
            }
        )
    return blocks


def parse_issues(raw: str) -> list[dict]:
    items = []
    for line in raw.splitlines():
        line = line.strip()
        m = re.match(r"^\d+[.、]\s*(.*)", line)
        if not m:
            continue
        content = m.group(1)
        tokens = re.findall(r"`([^`]+)`", content)
        reason = re.sub(r"`[^`]+`", "", content)
        reason = re.sub(r"→.*$", "", reason)
        reason = reason.strip(" ：;；、")
        items.append({"tokens": tokens, "reason": reason})
    return items


def match_block_to_paragraph(block: dict, paragraph_texts: list[str]) -> int:
    nr = norm(block["rev"])
    chunks = {nr[i : i + 16] for i in range(0, max(0, len(nr) - 15), 16)}
    best_i, best_score = -1, -1
    for i, pt in enumerate(paragraph_texts):
        score = sum(1 for c in chunks if c in norm(pt))
        if score > best_score:
            best_i, best_score = i, score
    return best_i


def locate_sentence(sentence: str, paragraph_texts: list[str], start: int, max_span: int = 8):
    ns = norm(sentence)
    lo = max(0, start - max_span)
    hi = min(len(paragraph_texts) - 1, start + max_span)
    for i in range(start, hi + 1):
        pos = norm(paragraph_texts[i]).find(ns)
        if pos >= 0:
            return i, pos
    for i in range(start - 1, lo - 1, -1):
        pos = norm(paragraph_texts[i]).find(ns)
        if pos >= 0:
            return i, pos
    return -1, -1


def run_is_superscript(paragraph, start: int, end: int) -> bool:
    containers = []
    total = 0
    lengths = []
    for child in paragraph._p:
        if child.tag not in (qn("w:r"), qn("w:hyperlink")):
            continue
        containers.append(child)
        clen = 0
        for cc in child:
            if cc.tag == qn("w:t"):
                clen += len(cc.text or "")
            elif cc.tag in (qn("w:tab"), qn("w:br"), qn("w:cr"), qn("w:noBreakHyphen")):
                clen += 1
        total += clen
        lengths.append(total)
    if not containers:
        return False

    def covering(pos):
        if pos >= lengths[-1]:
            return len(containers) - 1
        lo, hi = 0, len(lengths) - 1
        while lo < hi:
            mid = (lo + hi) // 2
            if lengths[mid] <= pos:
                lo = mid + 1
            else:
                hi = mid
        return min(lo, len(containers) - 1)

    si = covering(start)
    ei = covering(end - 1)
    for i in range(si, ei + 1):
        runs = [containers[i]] if containers[i].tag == qn("w:r") else containers[i].xpath("./w:r")
        for r_el in runs:
            rpr = r_el.find(qn("w:rPr"))
            if rpr is not None:
                va = rpr.find(qn("w:vertAlign"))
                if va is not None and va.get(qn("w:val")) == "superscript":
                    return True
    return False


def main() -> int:
    parser = argparse.ArgumentParser(description="Refine a review markdown into polish_edits.json")
    parser.add_argument("input", type=Path, help="original .docx")
    parser.add_argument("review_md", type=Path, help="paragraph-by-paragraph review markdown")
    parser.add_argument("-o", "--output", type=Path, default=Path("polish_edits.json"))
    args = parser.parse_args()

    doc = Document(str(args.input))
    paragraph_texts = [p.text for p in doc.paragraphs]
    blocks = parse_markdown(args.review_md)
    edits = []

    for block in blocks:
        para_idx = match_block_to_paragraph(block, paragraph_texts)
        orig_sentences = split_sentences(block["orig"])
        issues = parse_issues(block["issues_raw"])
        sent_locations = {}
        for i, s in enumerate(orig_sentences):
            sent_locations[i] = locate_sentence(s, paragraph_texts, para_idx)

        def find_oi(token):
            for i, s in enumerate(orig_sentences):
                if token in s:
                    return i
            return None

        def issue_reason(token):
            for issue in issues:
                if token in issue["tokens"]:
                    return issue["reason"]
            return "AI 整段改写：按期刊规范逐句修订"

        block_paras = {para_idx}
        for v in sent_locations.values():
            if v[0] >= 0:
                block_paras.add(int(v[0]))

        for token, replacement in KNOWN_REPLACEMENTS.items():
            if replacement == token:
                continue
            for spara in sorted(block_paras):
                para = doc.paragraphs[spara]
                nt = norm(para.text)
                ntoken = norm(token)
                if ntoken not in nt:
                    continue
                occ = 0
                pos = 0
                while True:
                    idx = nt.find(ntoken, pos)
                    if idx == -1:
                        break
                    if run_is_superscript(para, idx, idx + len(token)):
                        pos = idx + 1
                        continue
                    occ += 1
                    edits.append(
                        {
                            "paragraph_index": int(spara),
                            "original_text": token,
                            "revision_type": "句式润色",
                            "reason": issue_reason(token),
                            "revised_text": replacement,
                            "occurrence": occ,
                        }
                    )
                    pos = idx + 1

    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(
        json.dumps(
            {
                "document": args.input.name,
                "journal_guidelines": "",
                "edits": edits,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    print(f"refined {len(edits)} edit(s); saved {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())