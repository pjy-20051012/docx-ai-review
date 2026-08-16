# -*- coding: utf-8 -*-
"""Offline test suite for the docx-ai-review run isolator and comment injector."""

from __future__ import annotations

import json
import sys
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from lxml import etree

sys.path.insert(0, str(Path(__file__).parent))
from ai_review_to_comments import (
    COMMENT_AUTHOR,
    COMMENT_INITIALS,
    PolishBatch,
    PolishEdit,
    ParagraphRewrite,
    ReviewItem,
    apply_reviews,
    apply_paragraph_rewrite_tracked,
    apply_tracked_edits,
    convert_polish_edits_to_reviews,
    verify_comment_integrity,
    verify_tracked_integrity,
)


def build_fixture(path: Path) -> None:
    doc = Document()
    p0 = doc.add_paragraph()
    r1 = p0.add_run("引言")
    r1.bold = True
    r2 = p0.add_run("研究背景与目的")
    r3 = p0.add_run("显著影响")
    r3.italic = True
    p1 = doc.add_paragraph("本研究的显著性水平为 0.05，但显著性检验并未通过。")
    p2 = doc.add_paragraph("前一句包含超链接：")
    r_plain = p2.add_run("前半段")
    try:
        part = doc.part
        rel_id = part.relate_to(
            "https://example.com/source",
            RT.HYPERLINK,
            is_external=True,
        )
        hyperlink = p2._p.makeelement(qn("w:hyperlink"), {qn("r:id"): rel_id})
        r_hl = p2._p.makeelement(qn("w:r"), {})
        t_el = p2._p.makeelement(qn("w:t"), {})
        t_el.text = "权威文献"
        r_hl.append(t_el)
        hyperlink.append(r_hl)
        p2._p.append(hyperlink)
    except Exception:
        r_plain.text = "前半段（无超链接测试环境）"
    p3 = doc.add_paragraph("最后一个段落，包含中文全角标点：甲、乙；丙。")
    doc.save(str(path))


def reviews_json(path: Path) -> None:
    reviews = [
        {
            "paragraph_index": 0,
            "anchor_text": "研究背景与目的",
            "category": "用词精炼",
            "problem_analysis": "表述略冗长，可精简",
            "suggested_revision": "研究目的",
        },
        {
            "paragraph_index": 0,
            "anchor_text": "显著影响",
            "category": "语法修正",
            "problem_analysis": "缺少宾语",
            "suggested_revision": "产生显著影响",
        },
        {
            "paragraph_index": 1,
            "anchor_text": "显著性",
            "category": "逻辑表达",
            "problem_analysis": "前后结论冲突",
            "suggested_revision": "检验结果",
            "occurrence": 2,
        },
        {
            "paragraph_index": 3,
            "anchor_text": "中文全角标点",
            "category": "学术规范",
            "problem_analysis": "标点使用不规范",
            "suggested_revision": "中文全角标点；",
            "paragraph_revision": "最后一个段落的整段改写示例，用于验证完整段落修改会随批注一并写入。",
        },
    ]
    path.write_text(json.dumps({"reviews": reviews}, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> int:
    with tempfile.TemporaryDirectory(prefix="docx_ai_review_test_") as tmp:
        tmp = Path(tmp)
        fixture = tmp / "fixture.docx"
        output = tmp / "annotated.docx"
        reviews = tmp / "reviews.json"
        build_fixture(fixture)
        reviews_json(reviews)

        doc = Document(str(fixture))
        before = [p.text for p in doc.paragraphs]
        raw = json.loads(reviews.read_text(encoding="utf-8"))
        items = [ReviewItem.model_validate(x) for x in raw["reviews"]]
        apply_reviews(doc, items)
        doc.save(str(output))

        after_doc = Document(str(output))
        after = [p.text for p in after_doc.paragraphs]
        if before != after:
            raise AssertionError("paragraph text changed after comment injection")
        if len(after_doc.comments) != len(items):
            raise AssertionError(f"expected {len(items)} comments, got {len(after_doc.comments)}")
        for c in after_doc.comments:
            if c.author != COMMENT_AUTHOR or c.initials != COMMENT_INITIALS:
                raise AssertionError(f"unexpected author/initials: {c.author}/{c.initials}")
            if not c.text.startswith("【"):
                raise AssertionError(f"comment text format wrong: {c.text!r}")
        if not any("整段修改参考" in c.text for c in after_doc.comments):
            raise AssertionError("paragraph_revision was not included in any comment")

        problems = verify_comment_integrity(output)
        if problems:
            raise AssertionError("comment integrity problems: " + "; ".join(problems))

        # Check every commentRangeStart has a matching end with the same id in the same paragraph.
        with zipfile.ZipFile(output) as zf:
            doc_xml = zf.read("word/document.xml")
            comments_xml = zf.read("word/comments.xml")
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        root = etree.fromstring(doc_xml)
        starts = root.xpath("//w:commentRangeStart", namespaces=ns)
        ends = root.xpath("//w:commentRangeEnd", namespaces=ns)
        if len(starts) != len(ends) or len(starts) != len(items):
            raise AssertionError("commentRangeStart/End count mismatch")
        start_ids = [el.get(qn("w:id")) for el in starts]
        end_ids = [el.get(qn("w:id")) for el in ends]
        if sorted(start_ids) != sorted(end_ids):
            raise AssertionError("commentRangeStart ids do not match commentRangeEnd ids")
        comments_root = etree.fromstring(comments_xml)
        comment_ids = [el.get(qn("w:id")) for el in comments_root.xpath("//w:comment", namespaces=ns)]
        if sorted(comment_ids) != sorted(start_ids):
            raise AssertionError("comments.xml ids do not match document markers")

        # Convert workflow: a polishing skill's structured edit list -> comments.
        polish_json = tmp / "polish_edits.json"
        polish_json.write_text(
            json.dumps(
                {
                    "document": "fixture.docx",
                    "journal_guidelines": "test journal",
                    "edits": [
                        {
                            "paragraph_index": 0,
                            "original_text": "研究背景与目的",
                            "revision_type": "用词精炼",
                            "reason": "表述略冗长，可精简",
                            "revised_text": "研究目的",
                            "occurrence": 1,
                            "whole_paragraph_revision": "整段改写参考文本。",
                        }
                    ],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        polish_batch = PolishBatch.model_validate(
            json.loads(polish_json.read_text(encoding="utf-8"))
        )
        converted = convert_polish_edits_to_reviews(after_doc.paragraphs, polish_batch.edits)
        if len(converted) != 1 or converted[0].category != "用词精炼":
            raise AssertionError("convert_polish_edits_to_reviews failed")
        if converted[0].paragraph_revision != "整段改写参考文本。":
            raise AssertionError("whole_paragraph_revision not carried through conversion")

        # Tracked-changes workflow: replace a phrase, verify w:ins/w:del + comment.
        tracked_doc = Document(str(fixture))
        tracked_edits = [
            PolishEdit(
                paragraph_index=0,
                original_text="研究背景与目的",
                revision_type="用词精炼",
                reason="表述略冗长，可精简",
                revised_text="研究目的",
            )
        ]
        tracked_count = apply_tracked_edits(tracked_doc, tracked_edits)
        tracked_out = tmp / "tracked.docx"
        tracked_doc.save(str(tracked_out))
        if tracked_count != 1:
            raise AssertionError("tracked edit not applied")
        problems = verify_tracked_integrity(tracked_out)
        if problems:
            raise AssertionError("tracked integrity problems: " + "; ".join(problems))
        with zipfile.ZipFile(tracked_out) as zf:
            tracked_xml = zf.read("word/document.xml").decode("utf-8")
            tracked_comments = zf.read("word/comments.xml").decode("utf-8")
        if "<w:del " not in tracked_xml or "<w:ins " not in tracked_xml:
            raise AssertionError("tracked w:del/w:ins missing")
        if "修改理由" not in tracked_comments:
            raise AssertionError("rationale comment missing")

        # Full-paragraph rewrite workflow: sentence-level tracked changes.
        rewrite_doc = Document(str(fixture))
        rewrite_out = tmp / "rewrite.docx"
        p3_original = rewrite_doc.paragraphs[3].text
        revised_p3 = (
            "最后一个段落已经按期刊规范改写，包含中文全角标点；"
            "甲、乙；丙。新增一句用于验证插入。"
        )
        rewrite_count = apply_paragraph_rewrite_tracked(
            rewrite_doc,
            3,
            revised_p3,
            "AI 整段改写测试",
        )
        rewrite_doc.save(str(rewrite_out))
        if rewrite_count < 2:
            raise AssertionError("rewrite did not produce enough sentence changes")
        rewrite_problems = verify_tracked_integrity(rewrite_out)
        if rewrite_problems:
            raise AssertionError("rewrite integrity problems: " + "; ".join(rewrite_problems))
        with zipfile.ZipFile(rewrite_out) as zf:
            rewrite_xml = zf.read("word/document.xml").decode("utf-8")
        if "<w:del " not in rewrite_xml or "<w:ins " not in rewrite_xml:
            raise AssertionError("rewrite tracked markers missing")
        with zipfile.ZipFile(rewrite_out) as zf:
            rewrite_comments = zf.read("word/comments.xml").decode("utf-8")
        if "整段改写测试" not in rewrite_comments:
            raise AssertionError("rewrite rationale comment missing")
        if "删除" not in rewrite_xml and "w:delText" not in rewrite_xml:
            raise AssertionError("deleted original text not preserved as tracked deletion")

        print("ALL TESTS PASSED")
        print(f"  paragraphs preserved: {len(after)}")
        print(f"  comments injected: {len(after_doc.comments)}")
        print("  convert workflow: ok")
        print("  tracked-changes workflow: ok")
        print("  paragraph-rewrite workflow: ok")
        print(f"  output: {output}")
        return 0


if __name__ == "__main__":
    sys.exit(main())