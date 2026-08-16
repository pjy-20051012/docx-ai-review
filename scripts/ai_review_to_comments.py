# -*- coding: utf-8 -*-
"""Inject AI review comments into a .docx without changing the body text.

Commands:
  python ai_review_to_comments.py dump-text input.docx [--output text.json]
  python ai_review_to_comments.py apply input.docx reviews.json -o output.docx [--verify-only]
"""

from __future__ import annotations

import argparse
import copy
import datetime as dt
import difflib
import json
import re
import sys
import zipfile
from bisect import bisect_left
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.text.run import Run
from lxml import etree
from pydantic import BaseModel, Field, ValidationError

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CATEGORIES = ("学术规范", "语法修正", "逻辑表达", "用词精炼", "句式润色")
COMMENT_AUTHOR = "AI Reviewer"
COMMENT_INITIALS = "AI"


def _word_date() -> str:
    return dt.datetime.now(dt.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _make_ins_run(paragraph, text: str, base_run_el, rpr_el) -> object:
    """Create a tracked-insert w:ins element carrying revised text (caller inserts it)."""
    p_el = paragraph._p
    r_new = copy.deepcopy(base_run_el)
    # Strip existing text/break children so the new run carries only revised text.
    for child in list(r_new):
        if etree.QName(child.tag).localname != "rPr":
            r_new.remove(child)
    t_new = p_el.makeelement(qn("w:t"), {})
    t_new.text = text
    if text != text.strip():
        t_new.set(qn("xml:space"), "preserve")
    r_new.append(t_new)
    ins = p_el.makeelement(qn("w:ins"), {})
    ins.set(qn("w:id"), str(_next_revision_id(p_el)))
    ins.set(qn("w:author"), COMMENT_AUTHOR)
    ins.set(qn("w:date"), _word_date())
    ins.append(r_new)
    return ins


def _next_revision_id(paragraph_el) -> int:
    used = [
        int(x)
        for x in paragraph_el.xpath(
            ".//w:ins/@w:id | .//w:del/@w:id"
        )
        if x.isdigit()
    ]
    return max(used, default=0) + 1


def _mark_deleted_run(paragraph_el, r_el) -> None:
    """Wrap r_el in w:del, converting w:t to w:delText; the run keeps its own rPr."""
    del_el = paragraph_el.makeelement(qn("w:del"), {})
    del_el.set(qn("w:id"), str(_next_revision_id(paragraph_el)))
    del_el.set(qn("w:author"), COMMENT_AUTHOR)
    del_el.set(qn("w:date"), _word_date())
    for child in list(r_el):
        tag = etree.QName(child.tag).localname
        if tag == "t":
            child.tag = qn("w:delText")
    r_el.addprevious(del_el)
    del_el.append(r_el)


def _next_comment_anchor(paragraph, inserted_run_el, deleted_run_els):
    """Return runs to anchor the rationale comment to."""
    runs = []
    if inserted_run_el is not None:
        runs.append(Run(inserted_run_el, paragraph))
    elif deleted_run_els:
        runs.append(Run(deleted_run_els[0], paragraph))
    return runs


class ReviewItem(BaseModel):
    paragraph_index: int = Field(..., ge=0, description="0-based paragraph index")
    anchor_text: str = Field(..., min_length=1, description="exact substring in the paragraph")
    category: str = Field(..., description="one of " + "、".join(CATEGORIES))
    problem_analysis: str = Field(..., min_length=1, description="concise diagnosis, <= 50 chars")
    suggested_revision: str = Field(..., min_length=1, description="local revision for the anchored text")
    occurrence: int = Field(1, ge=1, description="1-based occurrence when the phrase repeats")
    paragraph_revision: str = Field(
        "",
        description="optional full rewritten paragraph text supplied by the reviewing model",
    )

    def model_post_init(self, __context) -> None:
        if self.category not in CATEGORIES:
            raise ValueError(f"category must be one of {CATEGORIES}, got {self.category!r}")
        if not self.anchor_text.strip():
            raise ValueError("anchor_text must contain non-whitespace characters")


class ReviewBatch(BaseModel):
    reviews: List[ReviewItem]


def _paragraph_text(paragraph) -> str:
    return paragraph.text


def _container_text(container) -> str:
    parts = []
    for child in container:
        tag = etree.QName(child.tag).localname
        if tag == "t":
            parts.append(child.text or "")
        elif tag in ("tab", "br", "cr", "noBreakHyphen"):
            parts.append("\t" if tag == "tab" else "\n")
    return "".join(parts)


def _char_map(paragraph_el):
    """Return (containers, cumulative_text_lengths) aligned with Paragraph.text."""
    containers = []
    lengths = []
    total = 0
    for child in paragraph_el:
        tag = etree.QName(child.tag).localname
        if tag in ("r", "hyperlink"):
            containers.append(child)
            total += len(_container_text(child))
            lengths.append(total)
    return containers, lengths


def _covering_index(offsets, pos: int) -> int:
    idx = bisect_left(offsets, pos + 1)
    if idx >= len(offsets):
        return len(offsets) - 1
    return idx


def _split_container(container, local_pos: int, paragraph_el) -> None:
    """Split container text at local_pos, inserting a deep-copied suffix container."""
    text = _container_text(container)
    if local_pos <= 0 or local_pos >= len(text):
        return
    suffix = copy.deepcopy(container)
    content = [c for c in suffix if etree.QName(c.tag).localname != "rPr"]
    seen = 0
    split_child = None
    prefix_children = []
    for child in content:
        tag = etree.QName(child.tag).localname
        if tag == "t":
            child_len = len(child.text or "")
        elif tag in ("tab", "br", "cr", "noBreakHyphen"):
            child_len = 1
        else:
            child_len = 0
        if child_len == 0:
            continue
        if local_pos < seen + child_len:
            if tag != "t":
                raise ValueError(
                    "anchor boundary falls inside a tab/line-break; choose an anchor that starts/ends on text"
                )
            rel = local_pos - seen
            if rel > 0:
                child.text = (child.text or "")[:rel]
                split_child = child
                suffix_text = copy.deepcopy(child)
                suffix_text.text = (copy.deepcopy(child).text or "")[rel:]
                child.addnext(suffix_text)
            else:
                prefix_children.append(child)
            break
        if seen + child_len == local_pos:
            prefix_children.append(child)
            break
        seen += child_len
        prefix_children.append(child)
    # The suffix keeps everything strictly after local_pos. Remove prefix elements.
    for child in prefix_children:
        if child.getparent() is suffix:
            suffix.remove(child)
    if split_child is not None and split_child.getparent() is suffix:
        suffix.remove(split_child)
    container.addnext(suffix)


def _split_at_global(paragraph_el, global_pos: int):
    containers, offsets = _char_map(paragraph_el)
    idx = _covering_index(offsets, global_pos)
    container = containers[idx]
    prev = offsets[idx - 1] if idx > 0 else 0
    local = global_pos - prev
    _split_container(container, local, paragraph_el)


def _isolate_runs(paragraph, start: int, end: int) -> List[Run]:
    """Split runs so [start, end) maps to exact Run objects within one paragraph."""
    if end <= start:
        raise ValueError(f"anchor range must be non-empty, got [{start}, {end})")
    p_el = paragraph._p
    _split_at_global(p_el, start)
    _split_at_global(p_el, end)
    containers, offsets = _char_map(p_el)
    si = _covering_index(offsets, start)
    ei = _covering_index(offsets, end - 1)
    runs = []
    for i in range(si, ei + 1):
        container = containers[i]
        if etree.QName(container.tag).localname == "r":
            runs.append(Run(container, paragraph))
        else:
            raise ValueError(
                "review anchor crosses a hyperlink boundary; choose a shorter anchor_text"
            )
    if not runs:
        raise ValueError("anchor container has no runs")
    return runs


def apply_reviews(doc: Document, reviews: List[ReviewItem]) -> int:
    paragraphs = doc.paragraphs
    count = 0
    for item in reviews:
        if item.paragraph_index >= len(paragraphs):
            raise IndexError(
                f"paragraph_index {item.paragraph_index} out of range (document has {len(paragraphs)} paragraphs)"
            )
        paragraph = paragraphs[item.paragraph_index]
        text = _paragraph_text(paragraph)
        anchor = item.anchor_text
        found = -1
        for _ in range(item.occurrence):
            found = text.find(anchor, found + 1)
            if found == -1:
                break
        if found == -1:
            raise ValueError(
                f"anchor_text {anchor!r} occurrence {item.occurrence} not found in paragraph {item.paragraph_index}"
            )
        runs = _isolate_runs(paragraph, found, found + len(anchor))
        comment_text = (
            f"【{item.category}】诊断：{item.problem_analysis}\n建议修改：{item.suggested_revision}"
        )
        if item.paragraph_revision:
            comment_text += "\n\n整段修改参考：\n" + item.paragraph_revision
        doc.add_comment(
            runs=runs,
            text=comment_text,
            author=COMMENT_AUTHOR,
            initials=COMMENT_INITIALS,
        )
        count += 1
    return count


class PolishEdit(BaseModel):
    """One atomic revision produced by a paper-polishing skill."""

    paragraph_index: int = Field(..., ge=0)
    original_text: str = Field(..., min_length=1, description="verbatim text from the manuscript")
    revision_type: str = Field(
        ...,
        description="one of 学术规范 / 语法修正 / 逻辑表达 / 用词精炼 / 句式润色",
    )
    reason: str = Field(..., min_length=1, description="concise reason, <= 60 chars")
    revised_text: str = Field(..., description="replacement for the anchored span")
    occurrence: int = Field(1, ge=1)
    whole_paragraph_revision: str = Field(
        "",
        description="optional full rewritten paragraph when the polish skill also rewrote it",
    )
    paragraph_revision: str = Field(
        "",
        description="alias for whole_paragraph_revision used by the rewrite command",
    )

    def model_post_init(self, __context) -> None:
        if self.revision_type not in CATEGORIES:
            raise ValueError(
                f"revision_type must be one of {CATEGORIES}, got {self.revision_type!r}"
            )


class PolishBatch(BaseModel):
    document: str = Field("", description="title or filename of the manuscript")
    journal_guidelines: str = Field("", description="journal name or short style notes")
    edits: List[PolishEdit]


class ParagraphRewrite(BaseModel):
    paragraph_index: int = Field(..., ge=0)
    original_text: str = Field("", description="optional verification copy of the original paragraph")
    paragraph_revision: str = Field(..., min_length=1)
    reason: str = Field("AI 整段改写：按期刊规范逐句修订", min_length=1)


class RewriteBatch(BaseModel):
    document: str = Field("")
    journal_guidelines: str = Field("")
    paragraphs: List[ParagraphRewrite]


def convert_polish_edits_to_reviews(paragraphs, edits: List[PolishEdit]) -> List[ReviewItem]:
    """Turn a polish skill's structured edit list into comment review items."""
    reviews: List[ReviewItem] = []
    for e in edits:
        if e.paragraph_index >= len(paragraphs):
            raise IndexError(
                f"paragraph_index {e.paragraph_index} out of range (document has {len(paragraphs)} paragraphs)"
            )
        text = paragraphs[e.paragraph_index].text
        if e.original_text not in text:
            raise ValueError(
                f"original_text {e.original_text!r} not found in paragraph {e.paragraph_index}"
            )
        reviews.append(
            ReviewItem(
                paragraph_index=e.paragraph_index,
                anchor_text=e.original_text,
                category=e.revision_type,
                problem_analysis=e.reason,
                suggested_revision=e.revised_text,
                occurrence=e.occurrence,
                paragraph_revision=e.whole_paragraph_revision,
            )
        )
    return reviews


def apply_tracked_edits(doc: Document, edits: List[PolishEdit]) -> int:
    """Apply a polishing edit list as Word tracked changes with rationale comments."""
    paragraphs = doc.paragraphs
    count = 0
    for e in edits:
        if e.paragraph_index >= len(paragraphs):
            raise IndexError(
                f"paragraph_index {e.paragraph_index} out of range (document has {len(paragraphs)} paragraphs)"
            )
        paragraph = paragraphs[e.paragraph_index]
        p_el = paragraph._p
        text = paragraph.text
        anchor = e.original_text
        if anchor and anchor not in text:
            raise ValueError(
                f"original_text {anchor!r} not found in paragraph {e.paragraph_index}"
            )
        if anchor == e.revised_text:
            continue

        if anchor:
            found = -1
            for _ in range(e.occurrence):
                found = text.find(anchor, found + 1)
                if found == -1:
                    raise ValueError(
                        f"original_text {anchor!r} occurrence {e.occurrence} not found"
                    )
            runs = _isolate_runs(paragraph, found, found + len(anchor))
            base_run_el = runs[0]._r
            rpr_el = base_run_el.find(qn("w:rPr"))
            deleted_els = []
            for run in runs:
                _mark_deleted_run(p_el, run._r)
                deleted_els.append(run._r)
            inserted_el = None
            if e.revised_text:
                ins_el = _make_ins_run(paragraph, e.revised_text, base_run_el, rpr_el)
                last_el = deleted_els[-1]
                while last_el.getparent() is not p_el:
                    last_el = last_el.getparent()
                last_el.addnext(ins_el)
                inserted_el = ins_el.find(qn("w:r"))
        else:
            # Pure insertion: append after the last text-bearing container in the paragraph.
            containers, offsets = _char_map(p_el)
            if not containers:
                raise ValueError("cannot insert into empty paragraph")
            insert_after = containers[-1]
            base_run_el = None
            for r_el in insert_after.xpath("./w:r"):
                base_run_el = r_el
                break
            if base_run_el is None:
                raise ValueError("insertion anchor run not found")
            ins_el = _make_ins_run(paragraph, e.revised_text, base_run_el, base_run_el.find(qn("w:rPr")))
            insert_after.addnext(ins_el)
            inserted_el = ins_el.find(qn("w:r"))

        anchor_runs = _next_comment_anchor(paragraph, inserted_el, deleted_els if anchor else [])
        if anchor_runs:
            comment_text = (
                f"【{e.revision_type}】修改理由：{e.reason}"
            )
            doc.add_comment(
                runs=anchor_runs,
                text=comment_text,
                author=COMMENT_AUTHOR,
                initials=COMMENT_INITIALS,
            )
        count += 1
    return count


def _split_sentences(text: str) -> List[str]:
    """Split mixed English/Chinese text into sentence units, normalized for matching."""
    text = _norm_spaces_hyphens(text)
    boundary_idx = []
    i = 0
    while i < len(text) - 1:
        ch = text[i]
        nxt = text[i + 1]
        if ch in ".!?" and (nxt.isspace() or nxt in "\"'(（"):
            # Require the next word to start with uppercase for English periods.
            if ch == ".":
                j = i + 1
                while j < len(text) and text[j].isspace():
                    j += 1
                if j >= len(text) or not text[j].isupper():
                    i += 1
                    continue
            boundary_idx.append(i + 1)
        elif ch in "。！？" or (ch == "；" and (nxt.isspace() or "\u4e00" <= nxt <= "\u9fff")):
            boundary_idx.append(i + 1)
        i += 1
    parts = []
    prev = 0
    for pos in boundary_idx:
        part = text[prev:pos].strip()
        if part:
            parts.append(part)
        prev = pos
    tail = text[prev:].strip()
    if tail:
        parts.append(tail)
    sentences = []
    for part in parts:
        stripped = part.strip()
        if stripped:
            sentences.append(stripped)
    return sentences


def _norm_spaces_hyphens(text: str) -> str:
    """Normalize Unicode spaces/hyphens to ASCII equivalents without changing length."""
    out = []
    for ch in text:
        if ch == "\u00a0" or ch == "\u202f" or ch == "\u2009" or ch == "\u200a" or ch == "\u2007":
            out.append(" ")
        elif ch in "\u2010\u2011\u2012\u2013\u2014\u2015":
            out.append("-")
        else:
            out.append(ch)
    return "".join(out)


def _insert_tracked_sentence(paragraph, text: str, after_el, base_run_el) -> object:
    """Insert text as a tracked-insert sentence after after_el."""
    ins_el = _make_ins_run(paragraph, text, base_run_el, base_run_el.find(qn("w:rPr")))
    after_el.addnext(ins_el)
    return ins_el.find(qn("w:r"))


def _delete_tracked_sentence(paragraph, sentence: str) -> None:
    """Delete an exact sentence with tracked deletion (must be called before other edits)."""
    p_el = paragraph._p
    text = paragraph.text
    norm_text = _norm_spaces_hyphens(text)
    norm_sentence = _norm_spaces_hyphens(sentence)
    start = norm_text.find(norm_sentence)
    if start == -1:
        raise ValueError(f"rewrite target sentence not found: {norm_sentence[:60]}...")
    runs = _isolate_runs(paragraph, start, start + len(sentence))
    for run in runs:
        _mark_deleted_run(p_el, run._r)


def apply_paragraph_rewrite_tracked(
    doc: Document,
    paragraph_index: int,
    revised_paragraph: str,
    reason: str = "AI 整段改写：按期刊规范逐句修订",
) -> int:
    """Apply a full paragraph rewrite as sentence-level tracked changes with comments."""
    paragraphs = doc.paragraphs
    if paragraph_index >= len(paragraphs):
        raise IndexError(f"paragraph_index {paragraph_index} out of range")
    paragraph = paragraphs[paragraph_index]
    original = paragraph.text
    norm_original = _norm_spaces_hyphens(original)
    orig_sentences = _split_sentences(original)
    new_sentences = _split_sentences(revised_paragraph)
    if not orig_sentences or not new_sentences:
        raise ValueError(f"cannot split paragraph {paragraph_index} into sentences")

    sm = difflib.SequenceMatcher(a=orig_sentences, b=new_sentences, autojunk=False)
    ops = []
    for op, o1, o2, n1, n2 in sm.get_opcodes():
        if op != "equal":
            ops.append((op, o1, o2, n1, n2))

    # Locate every sentence span in the ORIGINAL paragraph before any mutation.
    span_by_idx = {}
    cursor = 0
    for i, s in enumerate(orig_sentences):
        ns = _norm_spaces_hyphens(s)
        start = norm_original.find(ns, cursor)
        if start == -1:
            start = norm_original.find(ns)
        if start == -1:
            raise ValueError(f"rewrite target sentence not found: {ns[:60]}...")
        span_by_idx[i] = (start, start + len(s))
        cursor = start + len(ns)

    # Split all boundaries upfront, then capture the exact container elements per sentence.
    for i in sorted(span_by_idx):
        start, end = span_by_idx[i]
        _split_at_global(paragraph._p, start)
        _split_at_global(paragraph._p, end)
    containers, offsets = _char_map(paragraph._p)
    sentence_containers = {}
    for i, (start, end) in span_by_idx.items():
        si = _covering_index(offsets, start)
        ei = _covering_index(offsets, end - 1)
        sentence_containers[i] = (si, ei)

    def mark_containers_deleted(si, ei):
        for j in range(si, ei + 1):
            container = containers[j]
            if etree.QName(container.tag).localname == "r":
                _mark_deleted_run(paragraph._p, container)
            elif container.tag == qn("w:hyperlink"):
                for r_el in container.xpath("./w:r"):
                    _mark_deleted_run(paragraph._p, r_el)

    changed = 0
    # Process ops from the end backwards using saved element references.
    for op, o1, o2, n1, n2 in reversed(ops):
        if op == "insert":
            anchor_el = paragraph._p[-1]
            base_run_el = None
            for r_el in paragraph._p.iter(qn("w:r")):
                base_run_el = r_el
            if base_run_el is None:
                base_run_el = paragraph._p.makeelement(qn("w:r"), {})
            for s in new_sentences[n1:n2]:
                ins_el = _make_ins_run(paragraph, s, base_run_el, base_run_el.find(qn("w:rPr")))
                anchor_el.addnext(ins_el)
                anchor_el = ins_el
                doc.add_comment(
                    runs=[Run(ins_el.find(qn("w:r")), paragraph)],
                    text=f"【句式润色】修改理由：{reason}",
                    author=COMMENT_AUTHOR,
                    initials=COMMENT_INITIALS,
                )
                changed += 1
            continue

        if op == "delete":
            for i in range(o2 - 1, o1 - 1, -1):
                si, ei = sentence_containers[i]
                mark_containers_deleted(si, ei)
                changed += 1
            continue

        # replace: pair deleted original sentences with revised sentences in order.
        for k, j in enumerate(range(o2 - 1, o1 - 1, -1)):
            si, ei = sentence_containers[j]
            mark_containers_deleted(si, ei)
            last_container = containers[ei]
            while last_container.getparent() is not None and last_container.getparent() is not paragraph._p:
                last_container = last_container.getparent()
            revised = new_sentences[n2 - 1 - k]
            base_run_el = containers[si]
            if etree.QName(base_run_el.tag).localname != "r":
                base_run_el = None
                for r_el in paragraph._p.iter(qn("w:r")):
                    base_run_el = r_el
            ins_el = _make_ins_run(paragraph, revised, base_run_el, base_run_el.find(qn("w:rPr")))
            last_container.addnext(ins_el)
            doc.add_comment(
                runs=[Run(ins_el.find(qn("w:r")), paragraph)],
                text=f"【句式润色】修改理由：{reason}",
                author=COMMENT_AUTHOR,
                initials=COMMENT_INITIALS,
            )
            changed += 1
    return changed


def verify_tracked_integrity(docx_path: Path) -> List[str]:
    problems: List[str] = []
    with zipfile.ZipFile(docx_path) as zf:
        doc_xml = zf.read("word/document.xml")
        comments_xml = zf.read("word/comments.xml")
    ns = {"w": W_NS}
    root = etree.fromstring(doc_xml)
    revs = root.xpath("//w:ins | //w:del", namespaces=ns)
    for el in revs:
        if not el.get(qn("w:author")) or not el.get(qn("w:date")):
            problems.append(f"revision missing author/date ({etree.QName(el.tag).localname})")
    comments_root = etree.fromstring(comments_xml)
    comment_ids = {
        el.get(qn("w:id"))
        for el in comments_root.xpath("//w:comment", namespaces=ns)
        if el.get(qn("w:id")) is not None
    }
    for el in root.xpath("//w:commentRangeStart | //w:commentRangeEnd | //w:commentReference", namespaces=ns):
        cid = el.get(qn("w:id"))
        if cid not in comment_ids:
            problems.append(f"missing comment entity for id={cid}")
    return problems


def verify_comment_integrity(docx_path: Path) -> List[str]:
    problems: List[str] = []
    with zipfile.ZipFile(docx_path) as zf:
        names = zf.namelist()
        if "word/comments.xml" not in names:
            return ["missing word/comments.xml"]
        doc_xml = zf.read("word/document.xml")
        comments_xml = zf.read("word/comments.xml")
    doc_root = etree.fromstring(doc_xml)
    comments_root = etree.fromstring(comments_xml)
    ns = {"w": W_NS}
    comment_ids = {
        el.get(qn("w:id"))
        for el in comments_root.xpath("//w:comment", namespaces=ns)
        if el.get(qn("w:id")) is not None
    }
    refs = doc_root.xpath("//w:commentRangeStart | //w:commentRangeEnd | //w:commentReference", namespaces=ns)
    if not refs:
        problems.append("no comment markers in document.xml")
    for el in refs:
        cid = el.get(qn("w:id"))
        if cid not in comment_ids:
            problems.append(f"missing comment entity for id={cid} ({etree.QName(el.tag).localname})")
    return problems


def main(argv=None) -> int:
    parser = argparse.ArgumentParser(description="AI review to Word comments")
    sub = parser.add_subparsers(dest="command", required=True)

    p_dump = sub.add_parser("dump-text", help="dump paragraphs with indices")
    p_dump.add_argument("input", type=Path)
    p_dump.add_argument("--output", type=Path, default=None)

    p_apply = sub.add_parser("apply", help="apply reviews JSON to a docx")
    p_apply.add_argument("input", type=Path)
    p_apply.add_argument("reviews", type=Path)
    p_apply.add_argument("-o", "--output", type=Path, default=None)
    p_apply.add_argument("--verify-only", action="store_true")

    p_convert = sub.add_parser(
        "convert",
        help="convert a paper-polishing skill's structured edit list into comments",
    )
    p_convert.add_argument("input", type=Path)
    p_convert.add_argument("polish_edits", type=Path)
    p_convert.add_argument("-o", "--output", type=Path, required=True)

    p_tracked = sub.add_parser(
        "tracked",
        help="apply polish edits as Word tracked changes with rationale comments",
    )
    p_tracked.add_argument("input", type=Path)
    p_tracked.add_argument("polish_edits", type=Path)
    p_tracked.add_argument("-o", "--output", type=Path, required=True)

    p_rewrite = sub.add_parser(
        "rewrite",
        help="apply full paragraph rewrites as sentence-level tracked changes",
    )
    p_rewrite.add_argument("input", type=Path)
    p_rewrite.add_argument("rewrites", type=Path)
    p_rewrite.add_argument("-o", "--output", type=Path, required=True)

    p_verify = sub.add_parser("verify", help="verify comment markers in an annotated docx")
    p_verify.add_argument("input", type=Path)

    args = parser.parse_args(argv)

    if args.command == "dump-text":
        doc = Document(str(args.input))
        entries = [
            {"paragraph_index": i, "text": p.text, "char_count": len(p.text)}
            for i, p in enumerate(doc.paragraphs)
        ]
        payload = {"document": str(args.input), "paragraphs": entries}
        output = json.dumps(payload, ensure_ascii=False, indent=2)
        if args.output:
            args.output.write_text(output, encoding="utf-8")
        else:
            print(output)
        return 0

    if args.command == "apply":
        if args.verify_only:
            problems = verify_comment_integrity(args.input)
            if problems:
                print("VERIFY FAILED")
                for p in problems:
                    print(" -", p)
                return 1
            print(f"VERIFY OK: {args.input}")
            return 0
        if not args.output:
            parser.error("apply requires -o/--output")
        raw = json.loads(args.reviews.read_text(encoding="utf-8"))
        try:
            batch = ReviewBatch.model_validate(raw)
        except ValidationError as exc:
            print("Invalid reviews JSON:")
            print(exc)
            return 2
        doc = Document(str(args.input))
        count = apply_reviews(doc, batch.reviews)
        args.output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(args.output))
        problems = verify_comment_integrity(args.output)
        if problems:
            print("VERIFY FAILED after save")
            for p in problems:
                print(" -", p)
            return 3
        reopened = Document(str(args.output))
        print(
            f"OK: injected {count} comment(s); saved {args.output}; "
            f"reopened with {len(reopened.comments)} comment(s)."
        )
        return 0

    if args.command == "convert":
        raw = json.loads(args.polish_edits.read_text(encoding="utf-8"))
        try:
            batch = PolishBatch.model_validate(raw)
        except ValidationError as exc:
            print("Invalid polish edits JSON:")
            print(exc)
            return 2
        doc = Document(str(args.input))
        reviews = convert_polish_edits_to_reviews(doc.paragraphs, batch.edits)
        count = apply_reviews(doc, reviews)
        args.output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(args.output))
        problems = verify_comment_integrity(args.output)
        if problems:
            print("VERIFY FAILED after save")
            for p in problems:
                print(" -", p)
            return 3
        print(
            f"OK: converted {len(reviews)} polish edit(s) into {count} comment(s); saved {args.output}"
        )
        return 0

    if args.command == "tracked":
        raw = json.loads(args.polish_edits.read_text(encoding="utf-8"))
        try:
            batch = PolishBatch.model_validate(raw)
        except ValidationError as exc:
            print("Invalid polish edits JSON:")
            print(exc)
            return 2
        doc = Document(str(args.input))
        count = apply_tracked_edits(doc, batch.edits)
        args.output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(args.output))
        problems = verify_tracked_integrity(args.output)
        if problems:
            print("TRACKED VERIFY FAILED")
            for p in problems:
                print(" -", p)
            return 3
        print(
            f"OK: applied {count} tracked change(s) with comments; saved {args.output}"
        )
        return 0

    if args.command == "rewrite":
        raw = json.loads(args.rewrites.read_text(encoding="utf-8"))
        try:
            batch = RewriteBatch.model_validate(raw)
        except ValidationError as exc:
            print("Invalid rewrite JSON:")
            print(exc)
            return 2
        doc = Document(str(args.input))
        total = 0
        for item in batch.paragraphs:
            total += apply_paragraph_rewrite_tracked(
                doc,
                item.paragraph_index,
                item.paragraph_revision,
                item.reason or "AI 整段改写：按期刊规范逐句修订",
            )
        args.output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(args.output))
        problems = verify_tracked_integrity(args.output)
        if problems:
            print("REWRITE VERIFY FAILED")
            for p in problems:
                print(" -", p)
            return 3
        print(
            f"OK: applied {total} sentence-level tracked change(s); saved {args.output}"
        )
        return 0

    if args.command == "verify":
        problems = verify_comment_integrity(args.input)
        if problems:
            print("VERIFY FAILED")
            for p in problems:
                print(" -", p)
            return 1
        print(f"VERIFY OK: {args.input}")
        return 0

    return 1


if __name__ == "__main__":
    sys.exit(main())